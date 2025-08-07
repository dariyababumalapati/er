from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from utils import read_json

def render_table(doc, table_data):
    rows = len(table_data)
    cols = max(len(row) for row in table_data)

    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    occupied = [[False] * cols for _ in range(rows)]

    for r, row in enumerate(table_data):
        c = 0
        for cell in row:
            while c < cols and occupied[r][c]:
                c += 1
            if cell is None:
                continue

            colspan = cell.get("colspan", 1)
            rowspan = cell.get("rowspan", 1)

            for i in range(r, r + rowspan):
                for j in range(c, c + colspan):
                    occupied[i][j] = True

            tc = table.cell(r, c)
            tc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # ✅ Vertical center

            if rowspan > 1 or colspan > 1:
                tc = tc.merge(table.cell(r + rowspan - 1, c + colspan - 1))
                tc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # reapply after merge

            # Clear and set alignment
            p = tc.paragraphs[0]
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # ✅ Horizontal center

            # Rich text
            if "rich_text" in cell:
                for part in cell["rich_text"]:
                    run = p.add_run(part.get("text", ""))
                    run.bold = True  # Always bold

                    # Optional size control
                    if "size" in part:
                        run.font.size = Pt(part["size"])
                    else:
                        run.font.size = Pt(11)  # Default size

                    # Color control
                    if "color" in part:
                        color = part["color"].lower()
                        color_map = {
                            "red": RGBColor(255, 0, 0),
                            "green": RGBColor(0, 128, 0),
                            "blue": RGBColor(0, 112, 192),
                            "gray": RGBColor(128, 128, 128),
                            "black": RGBColor(0, 0, 0),
                        }
                        run.font.color.rgb = color_map.get(color, RGBColor(0, 0, 0))
            else:
                # Fallback plain text
                value = str(cell.get("text", cell.get("value", "")))
                run = p.add_run(value)
                run.bold = True
                run.font.size = Pt(11)  # Default fallback size

            c += colspan

    return table

if __name__ == "__main__":
    # Generate the table data first
    # table_data = generate_table_data(building_data, fj_weights)
    in_file = "data/vt4.json"


    table_data = read_json(in_file)


    # Create a new Word doc
    doc = Document()
    render_table(doc, table_data)

    # Save to output file
    doc.save("data/176C_table.docx")
